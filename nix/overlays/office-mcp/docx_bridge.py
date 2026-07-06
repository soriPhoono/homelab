#!/usr/bin/env python3
"""Bridge script for python-docx operations.

Called by the office-mcp JS server as:
  python3 docx_bridge.py <command> '<json_args>'

Outputs JSON to stdout.
"""
import sys
import json
import copy

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn


def read_structure(args):
    """Read detailed paragraph/style/run info from a .docx."""
    doc = Document(args["path"])
    paragraphs = []
    for i, para in enumerate(doc.paragraphs):
        runs = []
        for run in para.runs:
            runs.append({
                "text": run.text,
                "bold": run.bold,
                "italic": run.italic,
                "underline": run.underline,
                "font_name": run.font.name,
                "font_size": str(run.font.size) if run.font.size else None,
                "color": (
                    str(run.font.color.rgb)
                    if run.font.color and run.font.color.rgb
                    else None
                ),
            })
        paragraphs.append({
            "index": i,
            "text": para.text,
            "style": para.style.name,
            "alignment": str(para.alignment) if para.alignment is not None else None,
            "runs": runs,
        })

    tables = []
    for i, table in enumerate(doc.tables):
        rows = [[cell.text for cell in row.cells] for row in table.rows]
        tables.append({"index": i, "rows": rows})

    return {"paragraphs": paragraphs, "tables": tables}


def replace_text(args):
    """Find/replace text preserving per-run formatting."""
    doc = Document(args["path"])
    replacements = args["replacements"]
    count = 0

    def do_replace(paragraph):
        nonlocal count
        for old, new in replacements.items():
            if old not in paragraph.text:
                continue
            # Try run-level replacement first (preserves formatting)
            for run in paragraph.runs:
                if old in run.text:
                    run.text = run.text.replace(old, new)
                    count += 1
            # If text spans multiple runs, do a paragraph-level fallback
            if old in paragraph.text:
                full = paragraph.text
                full = full.replace(old, new)
                if paragraph.runs:
                    paragraph.runs[0].text = full
                    for r in paragraph.runs[1:]:
                        r.text = ""
                    count += 1

    for para in doc.paragraphs:
        do_replace(para)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    do_replace(para)

    output = args.get("output", args["path"])
    doc.save(output)
    return {"success": True, "replacements_made": count, "output": output}


def _make_paragraph(parent_element, insert_index, p_data):
    """Create a new w:p element and insert it at the given index."""
    from docx.text.paragraph import Paragraph
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    new_p_elem = parent_element.makeelement(qn("w:p"), {})
    parent_element.insert(insert_index, new_p_elem)
    para = Paragraph(new_p_elem, None)

    # Style
    if "heading" in p_data:
        level = p_data["heading"]
        para.style = f"Heading {level}"

    # Alignment
    align_map = {
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }
    if "align" in p_data and p_data["align"] in align_map:
        para.alignment = align_map[p_data["align"]]

    # Text run
    run = para.add_run(p_data.get("text", ""))
    if p_data.get("bold"):
        run.bold = True
    if p_data.get("italic"):
        run.italic = True
    if p_data.get("underline"):
        run.underline = True
    if p_data.get("size"):
        run.font.size = Pt(p_data["size"])
    if p_data.get("color"):
        color_hex = p_data["color"].lstrip("#")
        run.font.color.rgb = RGBColor.from_string(color_hex)
    if p_data.get("font"):
        run.font.name = p_data["font"]

    return new_p_elem


def insert_after(args):
    """Insert paragraphs after a paragraph containing specific text."""
    doc = Document(args["path"])
    after_text = args["afterText"]
    paragraphs = args["paragraphs"]
    inserted = 0

    body = doc.element.body
    children = list(body)

    for idx, child in enumerate(children):
        if child.tag == qn("w:p"):
            from docx.text.paragraph import Paragraph
            para = Paragraph(child, None)
            if after_text in para.text:
                insert_at = idx + 1
                for j, p_data in enumerate(paragraphs):
                    _make_paragraph(body, insert_at + j, p_data)
                    inserted += 1
                break

    output = args.get("output", args["path"])
    doc.save(output)
    return {"success": True, "inserted": inserted, "output": output}


def append_table(args):
    """Append a table to the end of a document."""
    doc = Document(args["path"])
    headers = args["headers"]
    rows = args["rows"]

    all_rows = [headers] + rows
    table = doc.add_table(rows=len(all_rows), cols=len(headers))
    table.style = "Table Grid"

    for ri, row_data in enumerate(all_rows):
        for ci, cell_text in enumerate(row_data):
            table.rows[ri].cells[ci].text = str(cell_text)
            # Bold the header row
            if ri == 0:
                for para in table.rows[ri].cells[ci].paragraphs:
                    for run in para.runs:
                        run.bold = True

    output = args.get("output", args["path"])
    doc.save(output)
    return {
        "success": True,
        "rows_added": len(all_rows),
        "cols": len(headers),
        "output": output,
    }


def delete_paragraph(args):
    """Delete paragraphs containing specific text."""
    doc = Document(args["path"])
    containing_text = args["containingText"]
    deleted = 0

    body = doc.element.body
    to_remove = []
    for child in list(body):
        if child.tag == qn("w:p"):
            from docx.text.paragraph import Paragraph
            para = Paragraph(child, None)
            if containing_text in para.text:
                to_remove.append(child)

    for elem in to_remove:
        body.remove(elem)
        deleted += 1

    output = args.get("output", args["path"])
    doc.save(output)
    return {"success": True, "deleted": deleted, "output": output}


COMMANDS = {
    "read_structure": read_structure,
    "replace_text": replace_text,
    "insert_after": insert_after,
    "append_table": append_table,
    "delete_paragraph": delete_paragraph,
}


def main():
    if len(sys.argv) < 3:
        print(json.dumps({"error": "Usage: docx_bridge.py <command> '<json_args>'"}))
        sys.exit(1)

    cmd = sys.argv[1]
    try:
        cmd_args = json.loads(sys.argv[2])
    except json.JSONDecodeError as e:
        print(json.dumps({"error": f"Invalid JSON args: {e}"}))
        sys.exit(1)

    if cmd not in COMMANDS:
        print(json.dumps({"error": f"Unknown command: {cmd}"}))
        sys.exit(1)

    try:
        result = COMMANDS[cmd](cmd_args)
        print(json.dumps(result))
    except Exception as e:
        print(json.dumps({"error": str(e), "command": cmd}))
        sys.exit(1)


if __name__ == "__main__":
    main()
